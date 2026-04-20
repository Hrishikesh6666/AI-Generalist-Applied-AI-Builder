from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime

# UrbanRoof brand colors
YELLOW = RGBColor(0xFF, 0xC0, 0x00)
DARK = RGBColor(0x33, 0x33, 0x33)

def add_heading_styled(doc, text, level=1):
    """Add a styled heading with brand colors"""
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.color.rgb = YELLOW if level == 1 else DARK
        run.font.bold = True
    return para

def create_ddr_word_doc(ai_output, inspection_images, thermal_images, 
                         property_name="Property", filename="DDR_Report.docx"):
    doc = Document()
    
    # ── Page Setup ────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # ── Cover Header ──────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DETAILED DIAGNOSTIC REPORT")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = YELLOW

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(f"Property: {property_name}")
    run2.font.size = Pt(12)
    run2.font.color.rgb = DARK

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated on: {datetime.now().strftime('%d %B %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = DARK
    
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()

    # ── Parse and write AI output ─────────────────────────────
    all_images = (
        [("Inspection", img) for img in inspection_images] + 
        [("Thermal", img) for img in thermal_images]
    )
    image_idx = 0  # track which image to insert next

    lines = ai_output.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Section headers (1. PROPERTY ISSUE SUMMARY etc)
        if (line[0].isdigit() and '.' in line[:3] and 
            line.split('.')[1].strip().isupper()):
            doc.add_page_break()
            add_heading_styled(doc, line, level=1)

        # Sub-headers (**bold text**)
        elif line.startswith('**') and line.endswith('**'):
            add_heading_styled(doc, line.replace('**', ''), level=2)

        # Image placeholder — insert actual image here
        elif '[IMAGE PLACEHOLDER:' in line:
            area_name = line.replace('[IMAGE PLACEHOLDER:', '').replace(']', '').strip()
            
            # Try to insert a relevant image
            if image_idx < len(all_images):
                source, img = all_images[image_idx]
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(io.BytesIO(img["data"]), width=Inches(4.5))
                    
                    caption = doc.add_paragraph(
                        f"[{source} Report - Page {img['page']}] {area_name}"
                    )
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.size = Pt(9)
                        run.font.italic = True
                    
                    image_idx += 1
                except Exception:
                    doc.add_paragraph(f"[Image for {area_name} - Not Available]")
            else:
                doc.add_paragraph(f"[Image: {area_name} - Not Available]")

        # Bullet points
        elif line.startswith('- ') or line.startswith('• '):
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(line[2:])

        # Normal paragraph
        else:
            doc.add_paragraph(line)

    # ── Remaining images in appendix ─────────────────────────
    if image_idx < len(all_images):
        doc.add_page_break()
        add_heading_styled(doc, "APPENDIX: SUPPORTING IMAGES", level=1)
        
        remaining = all_images[image_idx:]
        for source, img in remaining:
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(io.BytesIO(img["data"]), width=Inches(4.5))
                
                caption = doc.add_paragraph(
                    f"Source: {source} Report — Page {img['page']}"
                )
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
            except Exception:
                doc.add_paragraph(f"[Image from {source} Report, Page {img['page']} - Could not embed]")
    
    # ── Footer disclaimer ─────────────────────────────────────
    doc.add_page_break()
    doc.add_paragraph("─" * 80)
    disclaimer = doc.add_paragraph(
        "This report is generated by an AI-assisted system based on provided inspection "
        "and thermal documents. All findings should be verified by a qualified professional."
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].font.italic = True

    doc.save(filename)
    return filename